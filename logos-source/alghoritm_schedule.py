# pip install python-docx

import os
import re
import sqlite3
from datetime import datetime, timedelta
from docx import Document
from docx.oxml.ns import qn  # ✅ нужно для чтения цвета ячейки

DOCX_PATH = r"14.01.2026 1-2 ауысым  СОЛ апта.docx"

# --- Регексы ---
RE_DATE_ANY = re.compile(r"\b(\d{2}\.\d{2}\.\d{4})\b")
RE_PERIOD = re.compile(r"\b([1-9])\s*пара\b", re.IGNORECASE)

RE_TIME_ANY = re.compile(
    r"(\d{1,2})(?:[:\.]|)?([0-9⁰¹²³⁴⁵⁶⁷⁸⁹]{2})\s*[-–—]\s*(\d{1,2})(?:[:\.]|)?([0-9⁰¹²³⁴⁵⁶⁷⁸⁹]{2})"
)

RE_GROUP = re.compile(r"^[A-Za-zА-Яа-яЁёІіӘәӨөҮүҰұҚқҒғҺһ]+[0-9]{2}-[0-9]+[A-Za-zА-Яа-яЁёІіӘәӨөҮүҰұҚқҒғҺһ]*$")

SUP_MAP = str.maketrans({
    "⁰": "0", "¹": "1", "²": "2", "³": "3", "⁴": "4",
    "⁵": "5", "⁶": "6", "⁷": "7", "⁸": "8", "⁹": "9",
})

# ✅ Маппинг цвета -> подпись
COLOR_LABELS = {
    "D9D9D9": "Практика",
    "FFFF00": "Онлайн",
}


def clean(s: str) -> str:
    s = (s or "").replace("\u00a0", " ").strip()
    s = re.sub(r"[ \t]+", " ", s)
    s = "\n".join([x.strip() for x in s.splitlines() if x.strip()])
    return s


def supers_to_normal(s: str) -> str:
    return s.translate(SUP_MAP)


def get_cell_fill_hex(cell):
    """
    Возвращает HEX заливки ячейки (например 'FFFF00') или None если нет/auto.
    """
    tcPr = cell._tc.tcPr
    if tcPr is None:
        return None
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        return None
    fill = shd.get(qn("w:fill"))
    if not fill or fill.lower() == "auto":
        return None
    return fill.upper()


def label_by_fill(fill_hex: str | None) -> str:
    """
    Возвращает ' (Практика)' / ' (Онлайн)' или ''.
    """
    if not fill_hex:
        return ""
    label = COLOR_LABELS.get(fill_hex.upper())
    return f" ({label})" if label else ""


def find_schedule_date(doc: Document) -> str:
    for p in doc.paragraphs:
        m = RE_DATE_ANY.search(clean(p.text))
        if m:
            return m.group(1)

    for t in doc.tables:
        if not t.rows:
            continue
        header_text = " ".join(clean(c.text) for c in t.rows[0].cells)
        m = RE_DATE_ANY.search(header_text)
        if m:
            return m.group(1)

    raise RuntimeError("Не смог найти дату в документе (dd.mm.yyyy).")


def db_filename_for_date(ddmmyyyy: str) -> str:
    return f"schedule-{ddmmyyyy}.sqlite"


def cleanup_old_dbs(folder: str, older_than_days: int = 31):
    cutoff = datetime.now() - timedelta(days=older_than_days)
    for name in os.listdir(folder):
        if not (name.startswith("schedule-") and name.endswith(".sqlite")):
            continue
        m = re.match(r"^schedule-(\d{2}\.\d{2}\.\d{4})\.sqlite$", name)
        if not m:
            continue
        try:
            dt = datetime.strptime(m.group(1), "%d.%m.%Y")
        except ValueError:
            continue
        if dt < cutoff:
            try:
                os.remove(os.path.join(folder, name))
            except Exception:
                pass


def import_doc_file_and_tables(path: str):
    doc = Document(path)
    if not doc.tables:
        raise RuntimeError("В docx нет таблиц (doc.tables пуст).")
    t1 = doc.tables[0] if len(doc.tables) >= 1 else None
    t2 = doc.tables[1] if len(doc.tables) >= 2 else None
    return doc, t1, t2


def extract_times_from_header(table):
    if not table or not table.rows:
        return {}

    time_map = {}
    header_cells = table.rows[0].cells

    for cell in header_cells:
        txt = supers_to_normal(clean(cell.text))
        if not txt:
            continue

        p = RE_PERIOD.search(txt)
        t = RE_TIME_ANY.search(txt)
        if p and t:
            period = int(p.group(1))
            h1, m1, h2, m2 = t.groups()
            time_map[period] = f"{int(h1):02d}:{m1}-{int(h2):02d}:{m2}"

    if len(time_map) < 4:
        big = supers_to_normal(" ".join(supers_to_normal(clean(c.text)) for c in header_cells))
        for m in re.finditer(r"([1-9])\s*пара.*?" + RE_TIME_ANY.pattern, big, re.IGNORECASE):
            period = int(m.group(1))
            h1, m1, h2, m2 = m.group(2), m.group(3), m.group(4), m.group(5)
            time_map[period] = f"{int(h1):02d}:{m1}-{int(h2):02d}:{m2}"

    return time_map


def split_lines(cell_text: str):
    return [x.strip() for x in clean(cell_text).splitlines() if x.strip()]


def normalize_sportzaal(lesson_cell: str, group_cell: str):
    lessons = split_lines(lesson_cell)
    groups = split_lines(group_cell)

    if lessons and groups and len(lessons) == len(groups):
        return list(zip(lessons, groups))

    flat = split_lines(lesson_cell)
    pairs = []
    i = 0
    while i < len(flat) - 1:
        a = flat[i]
        b = flat[i + 1]
        if RE_GROUP.match(b):
            pairs.append((a, b))
            i += 2
        else:
            i += 1
    if pairs:
        return pairs

    if lessons:
        g = groups[0] if groups else ""
        return [(lessons[0], g)]
    return []


def create_schedule(table):
    """
    ✅ Теперь добавляем подпись к lesson по цвету ЯЧЕЙКИ (lesson-ячейки):
       D9D9D9 -> " (Практика)"
       FFFF00 -> " (Онлайн)"
       иначе ничего
    """
    schedule = {}

    for row in table.rows[1:]:
        # текст
        cells_text = [clean(c.text) for c in row.cells]
        if not cells_text:
            continue

        audience = cells_text[0].strip()
        if not audience:
            continue

        # сами cell-объекты (для чтения цвета)
        cells_obj = row.cells
        rest_text = cells_text[1:]
        pairs_mode = (len(rest_text) % 2 == 0 and len(rest_text) >= 2)

        if audience.lower().startswith("спорт зал"):
            out = []
            if pairs_mode:
                # идём по парам: (lesson_cell, group_cell)
                for i in range(0, len(rest_text), 2):
                    lesson_cell_obj = cells_obj[1 + i]
                    group_cell_obj = cells_obj[1 + i + 1]

                    fill = get_cell_fill_hex(lesson_cell_obj) or get_cell_fill_hex(group_cell_obj)
                    label = label_by_fill(fill)

                    # нормализуем спортзал
                    multi = normalize_sportzaal(rest_text[i], rest_text[i + 1])

                    # ✅ добавляем подпись ко всем урокам этой пары
                    multi_labeled = []
                    for lesson_text, group_text in multi:
                        lesson_text = (lesson_text.strip() + label).strip()
                        multi_labeled.append((lesson_text, group_text.strip()))
                    out.append(multi_labeled)
            else:
                for idx, cell_text in enumerate(rest_text):
                    cell_obj = cells_obj[1 + idx]
                    fill = get_cell_fill_hex(cell_obj)
                    label = label_by_fill(fill)

                    multi = normalize_sportzaal(cell_text, "")
                    multi_labeled = []
                    for lesson_text, group_text in multi:
                        lesson_text = (lesson_text.strip() + label).strip()
                        multi_labeled.append((lesson_text, group_text.strip()))
                    out.append(multi_labeled)

            schedule[audience] = out

        else:
            out = []
            if pairs_mode:
                for i in range(0, len(rest_text), 2):
                    lesson_cell_obj = cells_obj[1 + i]
                    group_cell_obj = cells_obj[1 + i + 1]

                    fill = get_cell_fill_hex(lesson_cell_obj) or get_cell_fill_hex(group_cell_obj)
                    label = label_by_fill(fill)

                    lesson = (rest_text[i].strip() + label).strip()
                    group = rest_text[i + 1].strip()
                    if lesson or group:
                        out.append((lesson, group))
            else:
                # fallback: одна ячейка = "предмет\nгруппа"
                for idx, cell_text in enumerate(rest_text):
                    cell_obj = cells_obj[1 + idx]
                    fill = get_cell_fill_hex(cell_obj)
                    label = label_by_fill(fill)

                    lines = split_lines(cell_text)
                    if len(lines) >= 2 and RE_GROUP.match(lines[-1]):
                        lesson = ("\n".join(lines[:-1]).strip() + label).strip()
                        out.append((lesson, lines[-1]))
                    elif lines:
                        lesson = (lines[0].strip() + label).strip()
                        out.append((lesson, ""))

            schedule[audience] = out

    return schedule


def init_db(db_path: str):
    conn = sqlite3.connect(db_path)
    conn.execute("""
        CREATE TABLE lessons (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            shift INTEGER NOT NULL,
            period INTEGER NOT NULL,
            time TEXT NOT NULL,
            audience TEXT NOT NULL,
            lesson TEXT NOT NULL,
            group_name TEXT NOT NULL
        )
    """)
    conn.execute("CREATE INDEX idx_lessons_group ON lessons(group_name)")
    conn.execute("CREATE INDEX idx_lessons_shift_period ON lessons(shift, period)")
    conn.commit()
    return conn


def save_schedule(conn, schedule: dict, time_map: dict, shift: int):
    for audience, periods in schedule.items():
        if audience.lower().startswith("спорт зал"):
            for period_idx, multi in enumerate(periods, start=1):
                time_str = time_map.get(period_idx, "")
                for lesson_text, group_text in multi:
                    conn.execute(
                        "INSERT INTO lessons(shift, period, time, audience, lesson, group_name) VALUES(?,?,?,?,?,?)",
                        (shift, period_idx, time_str, audience, lesson_text.strip(), group_text.strip())
                    )
        else:
            for period_idx, (lesson_text, group_text) in enumerate(periods, start=1):
                time_str = time_map.get(period_idx, "")
                conn.execute(
                    "INSERT INTO lessons(shift, period, time, audience, lesson, group_name) VALUES(?,?,?,?,?,?)",
                    (shift, period_idx, time_str, audience, lesson_text.strip(), group_text.strip())
                )
    conn.commit()


def print_schedule_for_group(conn, group_name: str):
    cur = conn.execute("""
        SELECT shift, period, time, audience, lesson
        FROM lessons
        WHERE group_name = ?
        ORDER BY shift ASC, period ASC, audience ASC
    """, (group_name,))
    rows = cur.fetchall()

    if not rows:
        print("Ничего не найдено для группы:", group_name)
        return

    for shift, period, time_str, audience, lesson in rows:
        print(f"Смена: {shift}")
        print(f"Каб: {audience}")
        print(f"Пара: {lesson}")
        print(f"Время: {time_str}")
        print("-" * 28)


def generate_schedule_db(docx_path: str, output_dir: str) -> str:
    doc, t1, t2 = import_doc_file_and_tables(docx_path)

    ddmmyyyy = find_schedule_date(doc)
    db_name = db_filename_for_date(ddmmyyyy)
    db_path = os.path.join(output_dir, db_name)

    if os.path.exists(db_path):
        os.remove(db_path)

    time_map_1 = extract_times_from_header(t1)
    schedule_1 = create_schedule(t1)

    time_map_2 = extract_times_from_header(t2) if t2 else {}
    schedule_2 = create_schedule(t2) if t2 else {}

    conn = init_db(db_path)
    save_schedule(conn, schedule_1, time_map_1, shift=1)
    if t2:
        save_schedule(conn, schedule_2, time_map_2, shift=2)
    conn.close()

    return db_path


# ---------------- RUN ----------------

script_dir = os.path.dirname(os.path.abspath(__file__))

if __name__ == "__main__":
    cleanup_old_dbs(script_dir, older_than_days=31)

    db_path = generate_schedule_db(DOCX_PATH, script_dir)

    conn = sqlite3.connect(db_path)
    GROUP = input().strip()
    print_schedule_for_group(conn, GROUP)
    conn.close()
